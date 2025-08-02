from fastapi import APIRouter, UploadFile, File, HTTPException
from utils.ai_client import GeminiClient
from pypdf import PdfReader
import io
import docx
from fastapi.responses import FileResponse
from docx import Document
import tempfile
import os

router = APIRouter()
ai_client = GeminiClient()

def read_docx(file: io.BytesIO) -> str:
    """Reads text from a DOCX file."""
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

@router.post("/generate-notes", tags=["notes"])
async def generate_notes_from_file(file: UploadFile = File(...)):
    try:
        file_extension = file.filename.split('.')[-1].lower()
        contents = await file.read()
        file_stream = io.BytesIO(contents)

        text = ""
        if file_extension == 'pdf':
            reader = PdfReader(file_stream)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif file_extension == 'docx':
            text = read_docx(file_stream)
        elif file_extension == 'txt':
            text = contents.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_extension}. Please upload a PDF, DOCX, or TXT file.")

        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the file. It might be empty or scanned.")

        notes_text = await ai_client.generate_notes_from_text(text)
        
        if not notes_text:
            raise HTTPException(status_code=502, detail="AI service failed to generate notes. Please try again.")

        # Create DOCX file from notes_text and save it
        doc = Document()
        for line in notes_text.splitlines():
            doc.add_paragraph(line)
        
        # Save DOCX to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        filename = f"AI_Notes_{file.filename.rsplit('.', 1)[0]}.docx"
        
        # Return both the notes text and the file path for download
        return {
            "notes": notes_text,
            "docx_path": tmp_path,
            "filename": filename
        }

    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"An unexpected error occurred in generate_notes_from_file: {e}")
        raise HTTPException(status_code=500, detail="An unexpected server error occurred.")

@router.post("/generate-notes-from-topic", tags=["notes"])
async def generate_notes_from_topic(request: dict):
    try:
        topic = request.get("topic", "")
        day = request.get("day", 1)
        tasks = request.get("tasks", [])
        
        if not topic:
            raise HTTPException(status_code=400, detail="Topic is required")
        
        # Create a comprehensive prompt for the AI
        prompt = f"""
        Generate comprehensive study notes for the topic: "{topic}"
        
        Day: {day}
        Tasks: {', '.join(tasks) if tasks else 'No specific tasks provided'}
        
        Please provide:
        1. Key concepts and definitions
        2. Important points to remember
        3. Examples and applications
        4. Study tips and strategies
        5. Common questions and answers
        6. Summary and key takeaways
        
        Make the notes comprehensive, well-structured, and easy to understand for students.
        """
        
        notes_text = await ai_client.generate_notes_from_text(prompt)
        
        if not notes_text:
            raise HTTPException(status_code=502, detail="AI service failed to generate notes. Please try again.")
        
        return {
            "notes": notes_text,
            "topic": topic,
            "day": day
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"An unexpected error occurred in generate_notes_from_topic: {e}")
        raise HTTPException(status_code=500, detail="An unexpected server error occurred.")

@router.get("/download-notes/{file_path:path}", tags=["notes"])
async def download_notes(file_path: str):
    try:
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        filename = os.path.basename(file_path)
        return FileResponse(file_path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"An unexpected error occurred in download_notes: {e}")
        raise HTTPException(status_code=500, detail="An unexpected server error occurred.")